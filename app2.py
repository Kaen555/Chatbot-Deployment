import streamlit as st
import pandas as pd
import os
from langchain_core.messages import HumanMessage, AIMessage
from workflow import app  # Import the compiled workflow app
from transformers import pipeline
import PyPDF2
import docx
import asyncio
from io import BytesIO
import re

# Initialize summarization model
summarizer = pipeline("summarization")

# Streamlit app for Legal Chatbot
st.set_page_config(page_title="Legal Chatbot", page_icon="⚖️", layout="wide")

st.title("⚖️ Legal Chatbot Assistant")
st.markdown("""
Welcome to the Legal Chatbot Assistant!  
Ask legal-related questions, schedule reminders, look up cases, or find nearby lawyers.  
You can also upload a legal document for summarization or assistance in filling out forms.
""")

# Sidebar for additional options
with st.sidebar:
    st.header("About the Chatbot")
    st.markdown("""
    - **Purpose**: Assist with legal queries and tasks.
    - **Features**:
      - Legal document analysis.
      - Case lookup.
      - Reminder scheduling.
      - Finding nearby lawyers.
      - Legal resources.
      - Document filling assistance.
    - **Note**: This bot is for informational purposes and not a substitute for professional legal advice.
    """)
    st.info("Supported Commands: Schedule reminders, look up cases, find nearby lawyers, schedule appointments, and more.")

# Document Filling Assistance Section
st.subheader("Document Filling Assistance for Legal Forms")
uploaded_form = st.file_uploader("Upload a legal form (DOCX only):", type=["docx"])

def extract_placeholders(text):
    """
    Extract placeholders in the format {{FieldName}} from a text.
    """
    return re.findall(r"\{\{(.*?)\}\}", text)

def read_docx(file):
    """
    Read and extract text from a DOCX file.
    """
    doc = docx.Document(file)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

def fill_docx_template(file, user_inputs):
    """
    Fill placeholders in a DOCX file with user inputs.
    """
    doc = docx.Document(file)
    output_buffer = BytesIO()
    for paragraph in doc.paragraphs:
        for field, value in user_inputs.items():
            if f"{{{{{field}}}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{{{field}}}}}", value)

    # Save the filled document to buffer
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer

if uploaded_form:
    form_text = read_docx(uploaded_form)
    placeholders = extract_placeholders(form_text)

    if placeholders:
        st.write("Placeholders detected in the document:")
        st.write(", ".join(placeholders))
        st.subheader("Fill in the Following Fields")
        user_inputs = {}
        for placeholder in placeholders:
            user_inputs[placeholder] = st.text_input(f"{placeholder}:")

        if st.button("Generate Filled Document"):
            if all(user_inputs.values()):
                filled_document = fill_docx_template(uploaded_form, user_inputs)
                st.success("Document filled successfully!")
                st.download_button(
                    label="Download Completed Document",
                    data=filled_document,
                    file_name="completed_form.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("Please fill in all fields before generating the document.")
    else:
        st.warning("No placeholders found in the document. Please upload a document with placeholders like {{Name}}.")

# File upload section for summarization
st.subheader("Upload a Legal Document for Summarization")
uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "txt"])

def extract_text_from_file(file):
    """
    Extracts text content from the uploaded file.
    Supports PDF, DOCX, and TXT formats.
    """
    if file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    elif file.type == "text/plain":
        return file.read().decode("utf-8")
    else:
        return None

if uploaded_file:
    document_text = extract_text_from_file(uploaded_file)

    if document_text:
        st.write("File Uploaded Successfully!")
        st.text_area("Extracted Document Content", document_text, height=200)

        # Summarize the extracted text
        if st.button("Summarize Document"):
            with st.spinner("Summarizing the document..."):
                try:
                    summary = summarizer(document_text, max_length=150, min_length=30, do_sample=False)
                    st.subheader("Document Summary")
                    st.write(summary[0]['summary_text'])
                except Exception as e:
                    st.error(f"Error during summarization: {e}")
    else:
        st.error("Unable to extract text from the uploaded file. Please try another format (PDF, DOCX, or TXT).")

# User input section for chatbot
st.subheader("Chat with the Legal Chatbot")
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# Language selection
language = st.selectbox("Choose the language for responses:", ["English", "German", "French", "Hindi"])
language_mapping = {
    "English": "en",
    "German": "de",
    "French": "fr",
    "Hindi": "hi"
}
selected_language = language_mapping[language]

user_input = st.text_input("Ask a legal-related question or give a command:", placeholder="Type your query here...", key="user_input")

if st.button("Ask"):
    if user_input.strip():
        input_messages = [HumanMessage(user_input)]
        config = {"configurable": {"thread_id": "abc1"}}
        response_placeholder = st.empty()

        st.session_state.chat_history.append(HumanMessage(content=user_input))

        async def display_response():
            accumulated_response = ""
            chunk_counter = 0

            async for chunk, metadata in app.astream(
                {"messages": input_messages, "language": selected_language},
                config,
                stream_mode="messages"
            ):
                if isinstance(chunk, AIMessage):
                    accumulated_response += chunk.content
                    chunk_counter += 1
                    response_placeholder.text_area(
                        "Response",
                        accumulated_response,
                        height=200,
                        key=f"response_area_{chunk_counter}"
                    )

        asyncio.run(display_response())
    else:
        st.warning("Please enter a valid query.")

# Display chat history
st.subheader("Chat History")
for message in st.session_state.chat_history:
    if isinstance(message, HumanMessage):
        st.markdown(f"**You:** {message.content}")
    elif isinstance(message, AIMessage):
        st.markdown(f"**Chatbot:** {message.content}")

# Feedback form section
st.subheader("Feedback Form")
st.write("Please provide your valuable feedback to improve the Legal Chatbot Assistant.")

FEEDBACK_FILE = "feedback_responses.csv"

def save_feedback_to_csv(feedback_data):
    if not os.path.exists(FEEDBACK_FILE):
        pd.DataFrame([feedback_data]).to_csv(FEEDBACK_FILE, index=False)
    else:
        pd.DataFrame([feedback_data]).to_csv(FEEDBACK_FILE, mode='a', header=False, index=False)


# Questionnaire File
QUESTIONNAIRE_FILE = "questionnaire_responses.csv"

def save_questionnaire_to_csv(data):
    """
    Save questionnaire responses to CSV.
    """
    if not os.path.exists(QUESTIONNAIRE_FILE):
        pd.DataFrame([data]).to_csv(QUESTIONNAIRE_FILE, index=False)
    else:
        pd.DataFrame([data]).to_csv(QUESTIONNAIRE_FILE, mode='a', header=False, index=False)

# Guided Intake Questionnaire Section
st.subheader("Guided Intake Questionnaire for Legal Issues")
st.write("Answer the following questions to help us understand your legal issue better.")

questionnaire_responses = {}
questionnaire_responses["Name"] = st.text_input("What is your name?")
questionnaire_responses["Contact"] = st.text_input("What is your contact information?")
questionnaire_responses["Issue Type"] = st.selectbox(
    "What type of legal issue are you facing?",
    ["Family Law", "Criminal Law", "Corporate Law", "Civil Disputes", "Other"]
)
questionnaire_responses["Urgency"] = st.radio(
    "How urgent is your legal issue?",
    ["Immediate", "Within a Week", "Not Urgent"]
)
questionnaire_responses["Details"] = st.text_area("Please provide a brief description of your legal issue.")

if st.button("Submit Questionnaire"):
    if all(questionnaire_responses.values()):  # Ensure all fields are filled
        save_questionnaire_to_csv(questionnaire_responses)
        st.success("Thank you for submitting the questionnaire. We will get back to you shortly!")
    else:
        st.warning("Please fill out all fields before submitting.")


feedback = {}
feedback["User Experience"] = st.selectbox("How was your overall experience?", ["Excellent", "Good", "Average", "Poor"])
feedback["Ease of Use"] = st.selectbox("Was the chatbot easy to use?", ["Yes", "No"])
feedback["Accuracy"] = st.selectbox("Was the response accurate?", ["Yes", "No"])
feedback["Comments"] = st.text_area("Any additional comments or suggestions?")

if st.button("Submit Feedback"):
    save_feedback_to_csv(feedback)
    st.success("Thank you for your feedback!")
    st.write("**Summary of Feedback:**")
    for key, value in feedback.items():
        st.write(f"**{key}**: {value}")

# Footer
st.markdown("---")
st.markdown("© 2024 Legal Chatbot Assistant | Built with ❤️ and 🤖")
